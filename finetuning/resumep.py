import random
from docx import Document
from datetime import date, timedelta


def create_random_cv():

    def random_date(start_date, end_date):
        time_between_dates = end_date - start_date
        random_days = random.randrange(time_between_dates.days)
        return start_date + timedelta(days=random_days)

    def generate_job_dates(number_of_jobs, graduation_date):
        job_dates = []
        start_date = graduation_date + timedelta(days=+30) 
        for _ in range(number_of_jobs):
            end_date = start_date + timedelta(days=random.randint(365, 365*3)) 
            job_dates.append((start_date, end_date))
            start_date = end_date + timedelta(days=random.randint(1, 30)) 
        return job_dates
    
    graduation_year = random.randint(2010, 2018)
    graduation_date = date(graduation_year, 6, 1)
    
    start_year = 2010
    end_year = 2023
    companies = [
        ('Google Inc.', 'worked on developing scalable machine learning models'),
        ('Microsoft Corp.', 'focused on cloud infrastructure optimization'),
        ('Amazon Inc.', 'engaged in e-commerce platform enhancement projects'),
        ('Facebook Inc.', 'contributed to social media analytics tools'),
        ('Tesla Inc.', 'participated in the design of battery management systems'),
        ('Apple Inc.', 'involved in the development of innovative user interfaces for mobile devices'),
        ('Netflix Inc.', 'contributed to the algorithm for personalized content recommendation'),
        ('Airbnb Inc.', 'worked on enhancing the platform’s global accommodation solutions'),
        ('SpaceX', 'assisted in the development of reusable rocket technology'),
        ('Uber Technologies Inc.', 'focused on optimizing routing algorithms for ride-sharing services'),
        ('LinkedIn Corp.', 'engaged in developing features to improve professional networking'),
        ('Adobe Systems Inc.', 'worked on cloud-based creative software services'),
        ('Twitter Inc.', 'contributed to the development of real-time data analysis tools'),
        ('Samsung Electronics Co.', 'involved in the development of next-generation semiconductor technology'),
        ('Intel Corp.', 'worked on optimizing high-performance computing architectures'),
        ('Sony Corp.', 'participated in the development of cutting-edge gaming technologies'),
        ('NVIDIA Corp.', 'focused on advancements in graphics processing units (GPUs) for gaming and professional markets'),
        ('IBM Corp.', 'engaged in projects related to quantum computing and AI research'),
        ('Huawei Technologies Co.', 'worked on developing solutions for 5G telecommunications networks'),
        ('Oracle Corp.', 'contributed to the enhancement of database management systems'),
    ]


    first_names = [
        'John', 'Jane', 'Alex', 'Emily', 'Michael', 'Sarah', 'Robert', 'Jessica',
        'David', 'Linda', 'Olivia', 'James', 'Isabella', 'Benjamin', 'Sophia',
        'Daniel', 'Emma', 'Matthew', 'Ava', 'Lucas'
    ]

    last_names = [
        'Smith', 'Johnson', 'Williams', 'Jones', 'Brown', 'Davis', 'Miller', 'Wilson',
        'Moore', 'Taylor', 'Anderson', 'Thomas', 'Jackson', 'White', 'Harris',
        'Martin', 'Thompson', 'Garcia', 'Martinez', 'Robinson'
    ]

    job_titles = [
        'Software Engineer', 'Systems Analyst', 'Web Developer', 'Network Administrator',
        'Database Administrator', 'Product Manager', 'Data Scientist', 'DevOps Engineer',
        'Security Analyst', 'User Experience Designer', 'Mobile Application Developer',
        'Cloud Solutions Architect', 'Front-end Developer', 'Back-end Developer',
        'Full Stack Developer', 'Quality Assurance Engineer', 'Machine Learning Engineer',
        'IT Project Manager', 'Software Architect', 'Technical Support Specialist'
    ]

    skills = [
        'Python', 'Java', 'C++', 'JavaScript', 'SQL', 'Ruby', 'PHP', 'C#', 'Swift',
        'Kotlin', 'HTML', 'CSS', 'React', 'Angular', 'Node.js', 'Django', 'Flask',
        'Spring', 'TensorFlow', 'PyTorch', 'AWS', 'Azure', 'Docker', 'Kubernetes',
        'Git', 'Linux', 'Machine Learning', 'Deep Learning', 'Data Analysis',
        'Agile Methodologies', 'Scrum', 'Project Management'
    ]

    languages = [
        'English', 'Spanish', 'French', 'German', 'Chinese', 'Russian', 'Arabic',
        'Portuguese', 'Italian', 'Japanese', 'Korean', 'Dutch', 'Swedish', 'Norwegian',
        'Danish', 'Finnish', 'Turkish', 'Greek', 'Hebrew', 'Hindi'
    ]

    certificates = [
        'Certified Ethical Hacker', 'Cisco Certified Network Associate', 'AWS Certified Solutions Architect',
        'Microsoft Certified: Azure Administrator Associate', 'Google Cloud Certified - Professional Cloud Architect',
        'Oracle Certified Professional Java SE Programmer', 'Certified ScrumMaster', 'CompTIA Security+',
        'Certified Information Systems Security Professional (CISSP)', 'Project Management Professional (PMP)',
        'Certified Data Professional', 'Certified AI & Machine Learning Professional', 'Certified Big Data Professional',
        'Certified in Risk and Information Systems Control (CRISC)', 'Certified Information Security Manager (CISM)',
        'Certified Cloud Security Professional (CCSP)', 'Salesforce Certified Administrator',
        'Certified Kubernetes Administrator', 'Certified Blockchain Expert', 'Certified IoT Professional'
    ]

    universities = [
        'Massachusetts Institute of Technology', 'Stanford University', 'Harvard University',
        'California Institute of Technology', 'University of Oxford', 'ETH Zurich',
        'University of Cambridge', 'Imperial College London', 'University of Chicago',
        'UCL (University College London)', 'National University of Singapore', 'Princeton University',
        'Nanyang Technological University, Singapore', 'EPFL', 'Tsinghua University',
        'University of California, Berkeley', 'Columbia University', 'University of Michigan',
        'Peking University', 'University of Edinburgh'
    ]

    degrees = [
        'B.Sc. in Computer Engineering', 'B.Sc. in Computer Science', 'B.Sc. in Software Engineering',
        'B.Sc. in Electrical Engineering and Computer Science', 'B.Sc. in Information Technology',
        'B.Eng. in Computer Engineering', 'B.Eng. in Software Engineering',
        'B.Sc. in Applied Mathematics and Computer Science'
    ]

    graduation_years = ['2019', '2020', '2021', '2022', '2023']

    for i in range(10):
        doc = Document()
      
        name = f"{random.choice(first_names)} {random.choice(last_names)}"
        doc.add_heading(name, level=1)

        
        doc.add_heading('Education', level=2)
        degree = random.choice(degrees)
        university = random.choice(universities)
        doc.add_paragraph(f"{degree}, {university}, Graduated {graduation_date.strftime('%b %Y')}")

        
        doc.add_heading('Job Experience', level=2)
        number_of_jobs = random.randint(1, 3)
        job_dates = generate_job_dates(number_of_jobs, graduation_date)
        for start_date, end_date in job_dates:
            company, work_description = random.choice(companies)
            job_title = random.choice(job_titles)
            experience = f"{job_title} at {company}, {start_date.strftime('%b %Y')} - {end_date.strftime('%b %Y')}"
            description = f"While at {company}, {work_description}."
            doc.add_paragraph(experience, style='ListBullet')
            doc.add_paragraph(description, style='BodyText')

       
        doc.add_heading('Skills', level=2)
        for skill in random.sample(skills, random.randint(1, 10)): 
            doc.add_paragraph(skill, style='ListBullet')

       

        doc.add_heading('Certificates', level=2)
        for certificate in random.sample(certificates, random.randint(1, 5)):  
            doc.add_paragraph(certificate, style='ListBullet')

        filename = f"C:\\Users\\ayses\\Desktop\\script\\CV_{i+1}_{name.replace(' ', '_')}.docx"
        doc.save(filename)
        print(f"CV saved: {filename}")

create_random_cv()
